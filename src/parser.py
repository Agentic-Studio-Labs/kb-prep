"""Document parser for DOCX and PDF files.

Extracts text, headings, and structure into ParsedDocument objects.
Also provides Markdown conversion for size reduction.
"""

import os
import re
from pathlib import Path
from typing import Optional

import fitz  # pymupdf
from docx import Document as DocxDocument

from .models import (
    DocumentMetadata,
    HeadingNode,
    Paragraph,
    ParsedDocument,
)

# Supported file extensions
SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".txt", ".md"}

# File size limits
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
WARN_FILE_SIZE = 25 * 1024 * 1024  # 25 MB


def discover_files(path: str, exclude_patterns: list[str] | None = None) -> list[str]:
    """Find all supported files in a path (file or directory).

    Args:
        path: File or directory path to scan.
        exclude_patterns: List of substrings to exclude from results
            (e.g. ["ORIGINAL DRAFTS", "OLD", "DO NOT USE"]).
    """
    p = Path(path)
    if p.is_file():
        if p.suffix.lower() in SUPPORTED_EXTENSIONS:
            return [str(p)]
        return []
    elif p.is_dir():
        files = []
        for ext in SUPPORTED_EXTENSIONS:
            files.extend(str(f) for f in p.rglob(f"*{ext}"))

        if exclude_patterns:
            files = [f for f in files if not any(pat in f for pat in exclude_patterns)]

        return sorted(files)
    return []


class DocumentParser:
    """Unified parser for DOCX, PDF, TXT, and MD files."""

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse a document and return structured representation."""
        p = Path(file_path)
        ext = p.suffix.lower()

        if ext == ".docx":
            return self._parse_docx(file_path)
        elif ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext in (".txt", ".md"):
            return self._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    # ------------------------------------------------------------------
    # DOCX parsing
    # ------------------------------------------------------------------

    def _parse_docx(self, file_path: str) -> ParsedDocument:
        doc = DocxDocument(file_path)
        paragraphs: list[Paragraph] = []
        idx = 0

        # Walk body elements in document order (paragraphs + tables)
        from docx.oxml.ns import qn

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Normal paragraph — use style-based heading detection
                para_obj = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para_obj = p
                        break
                if para_obj is None:
                    continue
                text = para_obj.text.strip()
                if not text:
                    continue
                level = self._docx_heading_level(para_obj.style.name)
                paragraphs.append(Paragraph(text=text, level=level, style=para_obj.style.name, index=idx))
                idx += 1

            elif tag == "tbl":
                # Table — extract cell text in row order with merged-cell dedup
                for tr in element.findall(qn("w:tr")):
                    seen_in_row: set[str] = set()
                    for tc in tr.findall(qn("w:tc")):
                        for p in tc.findall(qn("w:p")):
                            runs = p.findall(qn("w:r"))
                            text = "".join(
                                r.text for r in (run.find(qn("w:t")) for run in runs) if r is not None and r.text
                            )
                            text = text.strip()
                            if not text:
                                continue
                            # Deduplicate merged cells (same text in same row)
                            norm = " ".join(text.lower().split())
                            if norm in seen_in_row:
                                continue
                            seen_in_row.add(norm)
                            paragraphs.append(Paragraph(text=text, level=0, style="Table Cell", index=idx))
                            idx += 1

        # Extract metadata
        core = doc.core_properties
        file_size = os.path.getsize(file_path)
        metadata = DocumentMetadata(
            file_path=file_path,
            file_type="docx",
            file_size_bytes=file_size,
            title=core.title or None,
            author=core.author or None,
        )

        heading_tree = self._build_heading_tree(paragraphs)

        return ParsedDocument(
            metadata=metadata,
            paragraphs=paragraphs,
            heading_tree=heading_tree,
        )

    @staticmethod
    def _docx_heading_level(style_name: str) -> int:
        """Extract heading level from DOCX style name."""
        if not style_name:
            return 0
        match = re.match(r"Heading\s*(\d+)", style_name, re.IGNORECASE)
        if match:
            return int(match.group(1))
        if style_name.lower() == "title":
            return 1
        if style_name.lower() == "subtitle":
            return 2
        return 0

    # ------------------------------------------------------------------
    # PDF parsing
    # ------------------------------------------------------------------

    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        with fitz.open(file_path) as pdf:
            return self._parse_pdf_content(pdf, file_path)

    def _parse_pdf_content(self, pdf, file_path: str) -> ParsedDocument:
        paragraphs: list[Paragraph] = []
        idx = 0

        for page in pdf:
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

            for block in blocks:
                if block["type"] != 0:  # text block only
                    continue

                for line_group in block.get("lines", []):
                    spans = line_group.get("spans", [])
                    if not spans:
                        continue

                    text = "".join(s["text"] for s in spans).strip()
                    if not text:
                        continue

                    # Estimate heading level from font size
                    max_font_size = max(s["size"] for s in spans)
                    is_bold = any("bold" in s.get("font", "").lower() for s in spans)
                    level = self._estimate_heading_level(max_font_size, is_bold)

                    paragraphs.append(
                        Paragraph(
                            text=text,
                            level=level,
                            style=f"pdf-size-{max_font_size:.1f}",
                            index=idx,
                        )
                    )
                    idx += 1

        # Merge consecutive body paragraphs that are likely the same paragraph
        paragraphs = self._merge_pdf_paragraphs(paragraphs)

        file_size = os.path.getsize(file_path)
        metadata = DocumentMetadata(
            file_path=file_path,
            file_type="pdf",
            file_size_bytes=file_size,
            page_count=len(pdf),
            title=pdf.metadata.get("title") or None,
            author=pdf.metadata.get("author") or None,
        )

        heading_tree = self._build_heading_tree(paragraphs)

        return ParsedDocument(
            metadata=metadata,
            paragraphs=paragraphs,
            heading_tree=heading_tree,
        )

    @staticmethod
    def _estimate_heading_level(font_size: float, is_bold: bool) -> int:
        """Estimate heading level from PDF font characteristics."""
        if font_size >= 20:
            return 1
        elif font_size >= 16:
            return 2
        elif font_size >= 14 and is_bold:
            return 3
        elif font_size >= 13 and is_bold:
            return 4
        elif is_bold and font_size >= 11:
            return 5
        return 0

    @staticmethod
    def _merge_pdf_paragraphs(paragraphs: list[Paragraph]) -> list[Paragraph]:
        """Merge consecutive body-level lines into paragraphs."""
        if not paragraphs:
            return []

        merged: list[Paragraph] = []
        current: Optional[Paragraph] = None

        for p in paragraphs:
            if p.is_heading:
                if current:
                    merged.append(current)
                    current = None
                merged.append(p)
            elif current is None:
                current = Paragraph(
                    text=p.text,
                    level=0,
                    style=p.style,
                    index=p.index,
                )
            else:
                # Merge if the line looks like a continuation (no sentence-ending punctuation
                # at the end of current, or current is short)
                if not current.text.endswith((".", "!", "?", ":", ";")) or len(current.text.split()) < 15:
                    current = Paragraph(
                        text=current.text + " " + p.text,
                        level=0,
                        style=current.style,
                        index=current.index,
                    )
                else:
                    merged.append(current)
                    current = Paragraph(
                        text=p.text,
                        level=0,
                        style=p.style,
                        index=p.index,
                    )

        if current:
            merged.append(current)

        # Re-index with fresh Paragraph objects to avoid stale references
        return [Paragraph(text=p.text, level=p.level, style=p.style, index=i) for i, p in enumerate(merged)]

    # ------------------------------------------------------------------
    # Plain text / Markdown parsing
    # ------------------------------------------------------------------

    def _parse_text(self, file_path: str) -> ParsedDocument:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        paragraphs: list[Paragraph] = []
        is_md = file_path.lower().endswith(".md")

        for idx, block in enumerate(re.split(r"\n{2,}", content)):
            block = block.strip()
            if not block:
                continue

            level = 0
            style = "Normal"
            if is_md:
                md_match = re.match(r"^(#{1,6})\s+(.*)", block)
                if md_match:
                    level = len(md_match.group(1))
                    block = md_match.group(2)
                    style = f"Heading {level}"

            paragraphs.append(Paragraph(text=block, level=level, style=style, index=idx))

        file_size = os.path.getsize(file_path)
        metadata = DocumentMetadata(
            file_path=file_path,
            file_type="md" if is_md else "txt",
            file_size_bytes=file_size,
        )

        heading_tree = self._build_heading_tree(paragraphs)

        return ParsedDocument(
            metadata=metadata,
            paragraphs=paragraphs,
            heading_tree=heading_tree,
        )

    # ------------------------------------------------------------------
    # Heading tree builder
    # ------------------------------------------------------------------

    @staticmethod
    def _build_heading_tree(paragraphs: list[Paragraph]) -> list[HeadingNode]:
        """Build a hierarchical tree from flat heading list."""
        headings = [p for p in paragraphs if p.is_heading]
        if not headings:
            return []

        root_nodes: list[HeadingNode] = []
        stack: list[HeadingNode] = []

        for h in headings:
            node = HeadingNode(text=h.text, level=h.level, position=h.index)

            # Pop stack until we find a parent (lower level)
            while stack and stack[-1].level >= h.level:
                stack.pop()

            if stack:
                stack[-1].children.append(node)
            else:
                root_nodes.append(node)

            stack.append(node)

        return root_nodes


# ---------------------------------------------------------------------------
# Markdown conversion
# ---------------------------------------------------------------------------


def paragraphs_to_markdown(paragraphs: list[Paragraph]) -> str:
    """Render a list of Paragraphs as clean Markdown text."""
    lines: list[str] = []
    for para in paragraphs:
        if para.is_heading:
            prefix = "#" * para.level
            lines.append(f"{prefix} {para.text}")
        else:
            lines.append(para.text)
        lines.append("")  # blank line between paragraphs
    return "\n".join(lines).strip() + "\n"


def to_markdown(doc: ParsedDocument) -> str:
    """Convert a ParsedDocument to clean Markdown text.

    Useful for reducing file size before upload to a knowledge base,
    since Markdown strips formatting bloat from DOCX/PDF.
    """
    return paragraphs_to_markdown(doc.paragraphs)


def file_size_warning(file_size_bytes: int) -> Optional[str]:
    """Return a warning string if file size is concerning for RAG upload."""
    mb = file_size_bytes / (1024 * 1024)
    if file_size_bytes > MAX_FILE_SIZE:
        return f"CRITICAL: File is {mb:.1f} MB — exceeds 50 MB limit. Must reduce size."
    elif file_size_bytes > WARN_FILE_SIZE:
        return f"WARNING: File is {mb:.1f} MB — approaching 50 MB limit. Consider converting to Markdown."
    return None
