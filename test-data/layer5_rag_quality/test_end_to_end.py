"""End-to-end RAG round-trip evaluation.

Takes SQuAD benchmark text, writes it as .docx/.md files,
runs the full kb-prep pipeline, and measures retrieval quality
against a naive baseline.

This tests the complete path: parse → score → chunk → index → retrieve
"""

import sys
import tempfile
from pathlib import Path

import pytest

PROJECT_ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

pytestmark = [pytest.mark.layer5, pytest.mark.timeout(120)]

N_PARAGRAPHS = 50
N_DOCX = 10
TOP_K = 5


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _write_md_file(tmpdir: Path, idx: int, title: str, text: str) -> Path:
    """Write a SQuAD paragraph as a Markdown file."""
    safe_title = title[:60].replace("/", "-").replace("\\", "-")
    path = tmpdir / f"doc_{idx:03d}_{safe_title[:30]}.md"
    content = f"# {title}\n\n{text}\n"
    path.write_text(content, encoding="utf-8")
    return path


def _write_docx_file(tmpdir: Path, idx: int, title: str, text: str) -> Path:
    """Write a SQuAD paragraph as a DOCX file (heading + body sentences)."""
    from docx import Document

    path = tmpdir / f"doc_{idx:03d}_docx.docx"
    doc = Document()
    doc.add_heading(title, level=1)
    # Split text into sentences for more realistic paragraph structure
    import re

    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    # Group sentences into paragraphs of ~2-3 sentences each
    chunk_size = 3
    for i in range(0, len(sentences), chunk_size):
        group = " ".join(sentences[i : i + chunk_size])
        if group.strip():
            doc.add_paragraph(group)
    doc.save(str(path))
    return path


def _texttile_chunks(doc, metrics) -> list[str]:
    """Split a parsed document into chunks using TextTiling boundaries.

    Falls back to the full document as a single chunk if TextTiling
    doesn't find boundaries (common for short documents).
    """
    filename = doc.metadata.filename
    boundaries = []
    if metrics and filename in metrics:
        boundaries = metrics[filename].topic_boundaries

    para_texts = [p.text for p in doc.paragraphs if p.text.strip()]
    if not para_texts:
        return []

    if not boundaries:
        # Short document — use the full text as one chunk
        return [" ".join(para_texts)]

    # Split at boundaries
    chunks = []
    prev = 0
    for b in sorted(boundaries):
        if b > prev:
            chunk = " ".join(para_texts[prev:b])
            if chunk.strip():
                chunks.append(chunk)
        prev = b
    # Remainder
    if prev < len(para_texts):
        chunk = " ".join(para_texts[prev:])
        if chunk.strip():
            chunks.append(chunk)

    return chunks if chunks else [" ".join(para_texts)]


def _bm25_search(query: str, chunks: list[str], top_k: int = 5) -> list[int]:
    """Search chunks using BM25 and return top-k indices."""
    from src.corpus_analyzer import _bm25_score

    if not chunks:
        return []
    scores = _bm25_score(query, chunks)
    ranked = sorted(range(len(scores)), key=lambda i: -scores[i])
    return ranked[:top_k]


def _hit(answer: str, retrieved_chunks: list[str]) -> bool:
    """Check if the gold answer string appears in any retrieved chunk."""
    answer_lower = answer.lower()
    return any(answer_lower in chunk.lower() for chunk in retrieved_chunks)


# ---------------------------------------------------------------------------
# Engine pipeline
# ---------------------------------------------------------------------------


def run_engine_pipeline(
    tmpdir: Path,
    squad_sample: list[dict],
) -> tuple[list[str], list[dict[str, object]]]:
    """Parse, score, and chunk all files using the full kb-prep pipeline.

    Returns:
        (chunks, chunk_metadata) where chunk_metadata maps each chunk to
        its source paragraph_idx for diagnostics.
    """
    from src.corpus_analyzer import build_corpus_analysis
    from src.parser import DocumentParser
    from src.scorer import QualityScorer

    parser = DocumentParser()

    # Collect all file paths created in setup
    file_paths = sorted(tmpdir.glob("*.md")) + sorted(tmpdir.glob("*.docx"))
    assert file_paths, "No files found in temp dir"

    # Parse all documents
    parsed_docs = []
    for fp in file_paths:
        try:
            parsed_docs.append(parser.parse(str(fp)))
        except Exception as exc:  # noqa: BLE001
            print(f"  [warn] parse failed for {fp.name}: {exc}")

    assert parsed_docs, "All parses failed — pipeline broken"

    # Build corpus analysis (TF-IDF, TextTiling, entropy, self-retrieval)
    ca = build_corpus_analysis(parsed_docs)

    # Score each document (exercises scorer without crashing)
    scorer = QualityScorer(corpus_analysis=ca)
    for doc in parsed_docs:
        scorer.score(doc)  # Side-effect: raises on crash

    # Chunk using TextTiling boundaries
    all_chunks: list[str] = []
    all_meta: list[dict] = []

    for doc in parsed_docs:
        doc_chunks = _texttile_chunks(doc, ca.doc_metrics)
        for chunk in doc_chunks:
            all_chunks.append(chunk)
            all_meta.append({"file": doc.metadata.filename, "chunk": chunk[:80]})

    return all_chunks, all_meta


# ---------------------------------------------------------------------------
# Naive pipeline
# ---------------------------------------------------------------------------


def run_naive_pipeline(
    tmpdir: Path,
) -> list[str]:
    """Parse files and chunk with fixed-size sliding window (naive baseline)."""
    from layer5_rag_quality.naive_baseline import NaiveChunker
    from src.parser import DocumentParser

    parser = DocumentParser()
    chunker = NaiveChunker(chunk_size=200, overlap=30)

    file_paths = sorted(tmpdir.glob("*.md")) + sorted(tmpdir.glob("*.docx"))
    all_chunks: list[str] = []

    for fp in file_paths:
        try:
            doc = parser.parse(str(fp))
            text = doc.full_text
            if text.strip():
                all_chunks.extend(chunker.chunk(text))
        except Exception as exc:  # noqa: BLE001
            print(f"  [warn] naive parse failed for {fp.name}: {exc}")

    return all_chunks


# ---------------------------------------------------------------------------
# Main test
# ---------------------------------------------------------------------------


def test_end_to_end_rag_pipeline(squad_data):
    """Full pipeline: write synthetic docs → parse → chunk → index → retrieve.

    Compares TextTiling-chunked engine pipeline vs naive fixed-size chunking
    on SQuAD retrieval hit@5.
    """
    if len(squad_data) < N_PARAGRAPHS:
        pytest.skip(f"Need at least {N_PARAGRAPHS} SQuAD paragraphs, got {len(squad_data)}")

    sample = squad_data[:N_PARAGRAPHS]

    with tempfile.TemporaryDirectory() as tmpdir_str:
        tmpdir = Path(tmpdir_str)

        # ------------------------------------------------------------------
        # 1. Write synthetic documents
        # ------------------------------------------------------------------
        print(f"\n[setup] Writing {N_PARAGRAPHS} documents to {tmpdir}")

        for idx, record in enumerate(sample):
            text = record["text"]
            title = record["questions"][0]["question"] if record.get("questions") else f"Document {idx}"

            if idx < N_DOCX:
                _write_docx_file(tmpdir, idx, title, text)
            else:
                _write_md_file(tmpdir, idx, title, text)

        written = list(tmpdir.glob("*.md")) + list(tmpdir.glob("*.docx"))
        print(f"[setup] Wrote {len(written)} files ({N_DOCX} docx, {N_PARAGRAPHS - N_DOCX} md)")

        # ------------------------------------------------------------------
        # 2. Engine pipeline
        # ------------------------------------------------------------------
        print("[engine] Running full pipeline: parse → score → TextTile → index")
        engine_chunks, engine_meta = run_engine_pipeline(tmpdir, sample)
        print(f"[engine] Total chunks: {len(engine_chunks)}")

        # ------------------------------------------------------------------
        # 3. Naive pipeline
        # ------------------------------------------------------------------
        print("[naive] Running naive pipeline: parse → fixed-size chunk → index")
        naive_chunks = run_naive_pipeline(tmpdir)
        print(f"[naive] Total chunks: {len(naive_chunks)}")

        # ------------------------------------------------------------------
        # 4. Evaluate both pipelines
        # ------------------------------------------------------------------
        assert engine_chunks, "Engine produced no chunks — pipeline broken"
        assert naive_chunks, "Naive pipeline produced no chunks"

        engine_hits = 0
        naive_hits = 0
        total_questions = 0

        for record in sample:
            if not record.get("questions"):
                continue
            for qa in record["questions"][:1]:  # One question per paragraph
                query = qa["question"]
                answers = qa.get("answers", [])
                if not answers:
                    continue
                gold_answer = answers[0]
                total_questions += 1

                # Engine retrieval
                engine_top_indices = _bm25_search(query, engine_chunks, top_k=TOP_K)
                engine_retrieved = [engine_chunks[i] for i in engine_top_indices]
                if _hit(gold_answer, engine_retrieved):
                    engine_hits += 1

                # Naive retrieval
                naive_top_indices = _bm25_search(query, naive_chunks, top_k=TOP_K)
                naive_retrieved = [naive_chunks[i] for i in naive_top_indices]
                if _hit(gold_answer, naive_retrieved):
                    naive_hits += 1

        assert total_questions > 0, "No questions found in sample"

        engine_hit_rate = engine_hits / total_questions
        naive_hit_rate = naive_hits / total_questions

        # ------------------------------------------------------------------
        # 5. Print diagnostics
        # ------------------------------------------------------------------
        print(f"\n[results] Total questions evaluated: {total_questions}")
        print(f"[results] Engine hit@{TOP_K}: {engine_hit_rate:.3f} ({engine_hits}/{total_questions})")
        print(f"[results] Naive  hit@{TOP_K}: {naive_hit_rate:.3f} ({naive_hits}/{total_questions})")
        print(f"[results] Engine chunks: {len(engine_chunks)}, Naive chunks: {len(naive_chunks)}")

        # Sample retrieval diagnostic
        if sample and sample[0].get("questions"):
            sample_q = sample[0]["questions"][0]["question"]
            sample_ans = sample[0]["questions"][0]["answers"][0]
            sample_top = _bm25_search(sample_q, engine_chunks, top_k=3)
            print(f"\n[sample] Q: {sample_q!r}")
            print(f"[sample] Gold answer: {sample_ans!r}")
            for rank, i in enumerate(sample_top):
                hit_mark = "HIT" if sample_ans.lower() in engine_chunks[i].lower() else "miss"
                print(f"[sample] Rank {rank + 1} ({hit_mark}): {engine_chunks[i][:100]!r}...")

        # ------------------------------------------------------------------
        # 6. Assertions
        # ------------------------------------------------------------------
        assert engine_hit_rate > 0, (
            f"Engine hit@{TOP_K}=0 — pipeline produced results but nothing retrieved. "
            f"Chunks={len(engine_chunks)}, Questions={total_questions}"
        )

        # Engine should be competitive with naive (may lag on short docs where
        # TextTiling can't fire and falls back to full-doc chunks)
        if naive_hit_rate > 0:
            assert engine_hit_rate >= naive_hit_rate * 0.8, (
                f"Engine hit@{TOP_K}={engine_hit_rate:.3f} is more than 20% below "
                f"naive baseline={naive_hit_rate:.3f}. Engine chunking may be harming recall."
            )


# ---------------------------------------------------------------------------
# Smoke tests — individual pipeline stages
# ---------------------------------------------------------------------------


def test_pipeline_parse_stage_no_crash(squad_data):
    """Parsing synthetic .md and .docx files does not crash."""
    from src.parser import DocumentParser

    if len(squad_data) < 5:
        pytest.skip("Need SQuAD data")

    parser = DocumentParser()
    sample = squad_data[:5]

    with tempfile.TemporaryDirectory() as tmpdir_str:
        tmpdir = Path(tmpdir_str)

        for idx, record in enumerate(sample):
            title = record["questions"][0]["question"] if record.get("questions") else f"Doc {idx}"
            _write_md_file(tmpdir, idx, title, record["text"])

        for fp in tmpdir.glob("*.md"):
            doc = parser.parse(str(fp))
            assert doc.paragraphs, f"No paragraphs parsed from {fp.name}"
            assert doc.full_text.strip(), f"Empty full_text from {fp.name}"


def test_pipeline_corpus_analysis_no_crash(squad_data):
    """build_corpus_analysis does not crash on synthetic documents."""
    from src.corpus_analyzer import build_corpus_analysis
    from src.parser import DocumentParser

    if len(squad_data) < 5:
        pytest.skip("Need SQuAD data")

    parser = DocumentParser()
    sample = squad_data[:5]

    with tempfile.TemporaryDirectory() as tmpdir_str:
        tmpdir = Path(tmpdir_str)

        for idx, record in enumerate(sample):
            title = record["questions"][0]["question"] if record.get("questions") else f"Doc {idx}"
            _write_md_file(tmpdir, idx, title, record["text"])

        parsed_docs = [parser.parse(str(fp)) for fp in sorted(tmpdir.glob("*.md"))]
        ca = build_corpus_analysis(parsed_docs)

    assert ca.doc_metrics, "No doc_metrics produced"
    assert len(ca.doc_labels) == len(parsed_docs), "doc_labels count mismatch"


def test_pipeline_scorer_no_crash(squad_data):
    """QualityScorer.score does not crash on synthetic parsed documents."""
    from src.corpus_analyzer import build_corpus_analysis
    from src.parser import DocumentParser
    from src.scorer import QualityScorer

    if len(squad_data) < 5:
        pytest.skip("Need SQuAD data")

    parser = DocumentParser()
    sample = squad_data[:5]

    with tempfile.TemporaryDirectory() as tmpdir_str:
        tmpdir = Path(tmpdir_str)

        for idx, record in enumerate(sample):
            title = record["questions"][0]["question"] if record.get("questions") else f"Doc {idx}"
            _write_md_file(tmpdir, idx, title, record["text"])

        parsed_docs = [parser.parse(str(fp)) for fp in sorted(tmpdir.glob("*.md"))]
        ca = build_corpus_analysis(parsed_docs)
        scorer = QualityScorer(corpus_analysis=ca)

        for doc in parsed_docs:
            card = scorer.score(doc)
            assert card.overall_score >= 0
            assert card.overall_score <= 100
