"""Tests that documents survive the parse round-trip without losing content.

Sanity checks that the parsing layer doesn't silently drop content,
which would break any downstream retrieval.
"""

import sys
import tempfile
from pathlib import Path

import pytest

PROJECT_ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

pytestmark = [pytest.mark.layer5, pytest.mark.timeout(60)]


SAMPLE_PARAGRAPHS = [
    "The mitochondria is the powerhouse of the cell. It produces ATP through cellular respiration, which powers most biological processes in eukaryotic organisms.",
    "Photosynthesis converts sunlight into chemical energy stored as glucose. Chloroplasts in plant cells contain chlorophyll, which absorbs light in the red and blue spectrum.",
    "The water cycle describes continuous movement of water within Earth and its atmosphere. Evaporation, condensation, and precipitation are the three main stages of this cycle.",
]

SAMPLE_HEADING = "Biology and Earth Science Overview"


@pytest.fixture(scope="module")
def parser():
    from src.parser import DocumentParser

    return DocumentParser()


# ---------------------------------------------------------------------------
# DOCX round-trip
# ---------------------------------------------------------------------------


def test_docx_round_trip(parser):
    """Write text to .docx → parse → verify paragraphs match."""
    from docx import Document

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "sample.docx"
        doc = Document()
        doc.add_heading(SAMPLE_HEADING, level=1)
        for para_text in SAMPLE_PARAGRAPHS:
            doc.add_paragraph(para_text)
        doc.save(str(docx_path))

        parsed = parser.parse(str(docx_path))

    body_texts = [p.text for p in parsed.body_paragraphs]
    heading_texts = [p.text for p in parsed.headings]

    assert SAMPLE_HEADING in heading_texts, f"Heading lost after parse. Found headings: {heading_texts}"

    for expected in SAMPLE_PARAGRAPHS:
        assert any(expected in body for body in body_texts), (
            f"Paragraph lost after parse.\nExpected: {expected!r}\nFound: {body_texts}"
        )


# ---------------------------------------------------------------------------
# Markdown round-trip
# ---------------------------------------------------------------------------


def test_md_round_trip(parser):
    """Write text to .md → parse → verify paragraphs match."""
    with tempfile.TemporaryDirectory() as tmpdir:
        md_path = Path(tmpdir) / "sample.md"
        lines = [f"# {SAMPLE_HEADING}", ""]
        for para_text in SAMPLE_PARAGRAPHS:
            lines.append(para_text)
            lines.append("")
        md_path.write_text("\n".join(lines), encoding="utf-8")

        parsed = parser.parse(str(md_path))

    body_texts = [p.text for p in parsed.body_paragraphs]
    heading_texts = [p.text for p in parsed.headings]

    assert SAMPLE_HEADING in heading_texts, f"Heading lost after parse. Found headings: {heading_texts}"

    for expected in SAMPLE_PARAGRAPHS:
        assert any(expected in body for body in body_texts), (
            f"Paragraph lost after MD parse.\nExpected: {expected!r}\nFound: {body_texts}"
        )


# ---------------------------------------------------------------------------
# Plain text round-trip
# ---------------------------------------------------------------------------


def test_txt_round_trip(parser):
    """Write text to .txt → parse → verify content preserved."""
    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = Path(tmpdir) / "sample.txt"
        content = "\n\n".join(SAMPLE_PARAGRAPHS)
        txt_path.write_text(content, encoding="utf-8")

        parsed = parser.parse(str(txt_path))

    full = parsed.full_text
    for expected in SAMPLE_PARAGRAPHS:
        assert expected in full, f"Text lost in .txt parse.\nExpected: {expected!r}\nFound in full_text: {full!r}"


# ---------------------------------------------------------------------------
# Double-parse round-trip (parse → markdown → parse again)
# ---------------------------------------------------------------------------


def test_double_parse_round_trip(parser):
    """Parse → convert to markdown → parse again → verify content matches."""
    from src.parser import to_markdown

    with tempfile.TemporaryDirectory() as tmpdir:
        # Write the source document as .docx
        from docx import Document

        docx_path = Path(tmpdir) / "source.docx"
        doc = Document()
        doc.add_heading(SAMPLE_HEADING, level=1)
        for para_text in SAMPLE_PARAGRAPHS:
            doc.add_paragraph(para_text)
        doc.save(str(docx_path))

        # First parse
        parsed_1 = parser.parse(str(docx_path))
        md_content = to_markdown(parsed_1)

        assert md_content.strip(), "to_markdown returned empty string"

        # Write converted markdown and parse again
        md_path = Path(tmpdir) / "converted.md"
        md_path.write_text(md_content, encoding="utf-8")

        parsed_2 = parser.parse(str(md_path))

    body_texts_2 = [p.text for p in parsed_2.body_paragraphs]
    heading_texts_2 = [p.text for p in parsed_2.headings]

    assert SAMPLE_HEADING in heading_texts_2, f"Heading lost after double parse. Found: {heading_texts_2}"

    for expected in SAMPLE_PARAGRAPHS:
        assert any(expected in body for body in body_texts_2), (
            f"Paragraph lost after double parse.\nExpected: {expected!r}\nFound: {body_texts_2}"
        )


# ---------------------------------------------------------------------------
# No content loss on empty-ish documents
# ---------------------------------------------------------------------------


def test_parse_minimal_docx(parser):
    """A single-paragraph docx parses without crashing or dropping content."""
    from docx import Document

    text = "This is the only paragraph in this document and it must survive parsing."

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "minimal.docx"
        doc = Document()
        doc.add_paragraph(text)
        doc.save(str(docx_path))

        parsed = parser.parse(str(docx_path))

    assert text in parsed.full_text, "Single paragraph was dropped by parser"


def test_parse_minimal_md(parser):
    """A minimal markdown document parses correctly."""
    with tempfile.TemporaryDirectory() as tmpdir:
        md_path = Path(tmpdir) / "minimal.md"
        md_path.write_text("# Title\n\nBody content here.\n", encoding="utf-8")

        parsed = parser.parse(str(md_path))

    assert any(p.level == 1 for p in parsed.paragraphs), "Heading level 1 not found"
    assert any("Body content here" in p.text for p in parsed.body_paragraphs), "Body paragraph missing"
