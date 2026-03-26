"""Tests for DOCX table extraction in the parser."""

import os
import tempfile

from docx import Document

from src.parser import DocumentParser
from src.scorer import QualityScorer


def _create_table_heavy_docx(path: str):
    """Create a DOCX that mimics a lesson plan with tables (like the real lesson docs)."""
    doc = Document()
    doc.add_heading("Financial Literacy Lesson 2", level=1)
    doc.add_paragraph("Grades 4-5")

    # Lesson overview table (2 columns: label | content)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Lesson Title"
    table.cell(0, 1).text = "How Can SMART Goals Help You Achieve Your Dreams?"
    table.cell(1, 0).text = "Lesson Overview"
    table.cell(1, 1).text = (
        "Students will learn to set Specific Measurable Achievable Relevant "
        "and Time-bound goals to plan their financial future. This lesson "
        "introduces goal-setting frameworks through interactive activities."
    )
    table.cell(2, 0).text = "Materials"
    table.cell(2, 1).text = "Handout A, Handout B, anchor chart paper, markers"

    # Activities table
    doc.add_heading("Activities", level=2)
    activities = doc.add_table(rows=2, cols=2)
    activities.cell(0, 0).text = "Activity 1"
    activities.cell(0, 1).text = "Students brainstorm personal goals and categorize them as short-term or long-term."
    activities.cell(1, 0).text = "Activity 2"
    activities.cell(1, 1).text = "Students convert one goal into a SMART goal using the planning template."

    doc.save(path)


def test_docx_parser_extracts_table_cell_text():
    """Parser should extract text from table cells, not just paragraphs."""
    parser = DocumentParser()
    with tempfile.TemporaryDirectory() as td:
        path = os.path.join(td, "lesson.docx")
        _create_table_heavy_docx(path)
        doc = parser.parse(path)

        full_text = doc.full_text.lower()
        assert "smart goals" in full_text
        assert "lesson overview" in full_text
        assert "brainstorm personal goals" in full_text
        assert "planning template" in full_text


def test_docx_parser_preserves_body_order_with_tables():
    """Paragraphs and table content should appear in document order."""
    parser = DocumentParser()
    with tempfile.TemporaryDirectory() as td:
        path = os.path.join(td, "lesson.docx")
        _create_table_heavy_docx(path)
        doc = parser.parse(path)

        texts = [p.text for p in doc.paragraphs]
        # Title paragraph comes before table content
        title_idx = next(i for i, t in enumerate(texts) if "Financial Literacy" in t)
        overview_idx = next(i for i, t in enumerate(texts) if "Lesson Overview" in t)
        activity_idx = next(i for i, t in enumerate(texts) if "brainstorm personal goals" in t)
        assert title_idx < overview_idx < activity_idx


def test_docx_parser_deduplicates_merged_cells():
    """Merged cells should not produce duplicate text."""
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    # Simulate merged cell: same text in adjacent cells
    table.cell(0, 0).text = "Header Text"
    table.cell(0, 1).text = "Header Text"
    table.cell(0, 2).text = "Different Text"

    parser = DocumentParser()
    with tempfile.TemporaryDirectory() as td:
        path = os.path.join(td, "merged.docx")
        doc.save(path)
        parsed = parser.parse(path)

        header_count = sum(1 for p in parsed.paragraphs if p.text == "Header Text")
        assert header_count == 1, f"Expected 1 'Header Text', got {header_count}"


def test_low_parse_fidelity_flag_for_sparse_docx():
    """Scorer should flag when extracted text is suspiciously sparse for the file size."""
    # Create a minimal DOCX (will be small on disk, but we fake the file_size_bytes)
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Short.")

    parser = DocumentParser()
    with tempfile.TemporaryDirectory() as td:
        path = os.path.join(td, "sparse.docx")
        doc.save(path)
        parsed = parser.parse(path)
        # Override file size to simulate a large file with little extracted content
        parsed.metadata.file_size_bytes = 50_000

        scorer = QualityScorer()
        card = scorer.score(parsed)

        structure_issues = [i for r in card.results if r.category == "structure" for i in r.issues]
        fidelity_issues = [i for i in structure_issues if "parse fidelity" in i.message.lower()]
        assert len(fidelity_issues) == 1, f"Expected 1 fidelity warning, got {len(fidelity_issues)}"
