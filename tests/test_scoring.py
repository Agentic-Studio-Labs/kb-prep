#!/usr/bin/env python3
"""Test the scoring pipeline with a synthetic DOCX document.

Creates a lesson plan with intentional RAG quality issues,
then scores it and verifies detection.
"""

import os
import sys
import tempfile
from pathlib import Path

# Add parent to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document

from parser import DocumentParser
from scorer import QualityScorer


def _create_test_docx(path: str):
    """Create a DOCX with intentional quality issues."""
    doc = Document()

    # Title
    doc.add_heading("Lesson Plan", level=1)

    # Issue: generic heading
    doc.add_heading("Content", level=2)
    doc.add_paragraph(
        "This lesson covers the basics of adding fractions with unlike denominators. "
        "Students will learn to find common denominators and simplify their answers. "
        "The lesson is designed for fifth-grade students who have mastered basic fraction concepts."
    )

    # Issue: dangling reference
    doc.add_paragraph(
        "As mentioned above, students should already understand basic fractions. "
        "The teacher should review the steps outlined in the previous section before "
        "moving on to the exercises below."
    )

    # Issue: heading level jump (H2 → H4)
    doc.add_heading("Practice Problems", level=4)
    doc.add_paragraph("Complete the following exercises to reinforce understanding of fraction addition.")

    # Issue: very long paragraph (>300 words)
    long_text = (
        "Fraction addition is a fundamental mathematical skill that students encounter "
        "in elementary school. When adding fractions with unlike denominators, the first "
        "step is to find a common denominator. This is typically the least common multiple "
        "(LCM) of the two denominators. For example, when adding 1/3 and 1/4, the LCM of "
        "3 and 4 is 12. So we convert 1/3 to 4/12 and 1/4 to 3/12. Then we can add the "
        "numerators: 4/12 + 3/12 = 7/12. It is important that students understand why we "
        "need common denominators and not just memorize the procedure. Teachers should use "
        "visual aids such as fraction bars or pie charts to demonstrate how fractions with "
        "different denominators represent different-sized pieces. Students often struggle "
        "with this concept because they try to add numerators and denominators separately. "
        "For instance, they might incorrectly calculate 1/3 + 1/4 as 2/7, which shows a "
        "fundamental misunderstanding of what fractions represent. To address this, teachers "
        "can have students use manipulatives to physically combine fractional parts. Another "
        "common error is forgetting to simplify the final answer. After adding fractions, "
        "students should always check if the resulting fraction can be reduced. For example, "
        "2/4 should be simplified to 1/2. Teachers can reinforce this by asking students to "
        "always express their answers in simplest form. Additionally, mixed numbers present "
        "another challenge. When the sum of two fractions results in an improper fraction, "
        "students need to convert it to a mixed number. For example, 3/4 + 3/4 = 6/4, which "
        "should be expressed as 1 and 2/4, or simplified further to 1 and 1/2. Practice is "
        "essential for mastering fraction addition, and teachers should provide a variety of "
        "problems with increasing difficulty levels to ensure deep understanding of the concept."
    )
    doc.add_paragraph(long_text)

    # Issue: undefined acronym
    doc.add_heading("Assessment", level=2)
    doc.add_paragraph(
        "Use the IEP goals to track student progress. Students meeting CCSS benchmarks "
        "should advance to the next unit. CCSS alignment ensures consistent standards."
    )

    # Issue: very short paragraph
    doc.add_paragraph("Review answers.")

    # Another generic heading
    doc.add_heading("Notes", level=2)
    doc.add_paragraph("Teachers should adapt the lesson based on student needs and available materials.")

    doc.save(path)


def test_scoring_detects_expected_issues():
    """Score a synthetic document and verify expected issue categories are detected."""
    with tempfile.TemporaryDirectory() as tmpdir:
        # Create test file with a generic filename (another issue)
        test_path = os.path.join(tmpdir, "doc-v2.docx")
        _create_test_docx(test_path)

        # Parse
        parser = DocumentParser()
        doc = parser.parse(test_path)

        assert len(doc.paragraphs) > 0
        assert len(doc.headings) > 0

        # Score
        scorer = QualityScorer()
        card = scorer.score(doc)

        assert card.overall_score > 0
        assert len(card.all_issues) > 0

        # Validate that we caught the expected issues
        categories_with_issues = {i.category for i in card.all_issues}

        expected = {
            "self_containment",  # "as mentioned above", "previous section", "below"
            "heading_quality",  # "Content", "Notes" generic; level jump
            "paragraph_length",  # long paragraph + short paragraph
            "filename_quality",  # "doc-v2"
        }

        missing = expected - categories_with_issues
        assert not missing, f"Scorer missed these categories: {missing}"


def test_score_breakdown():
    """Verify score categories and weights are reasonable."""
    with tempfile.TemporaryDirectory() as tmpdir:
        test_path = os.path.join(tmpdir, "doc-v2.docx")
        _create_test_docx(test_path)

        parser = DocumentParser()
        doc = parser.parse(test_path)
        scorer = QualityScorer()
        card = scorer.score(doc)

        # All 8 base categories should be present
        categories = {r.category for r in card.results}
        assert "self_containment" in categories
        assert "heading_quality" in categories
        assert "paragraph_length" in categories
        assert "file_focus" in categories
        assert "filename_quality" in categories
        assert "acronym_definitions" in categories
        assert "structure" in categories
        assert "file_size" in categories

        # Weights of non-zero criteria should sum to ~1.0
        total_weight = sum(r.weight for r in card.results if r.weight > 0)
        assert abs(total_weight - 1.0) < 0.01, f"Weights sum to {total_weight}, expected ~1.0"


def test_new_scoring_criteria_present():
    """New criteria (readability, retrieval_aware) appear in results."""
    from corpus_analyzer import build_corpus_analysis

    with tempfile.TemporaryDirectory() as tmpdir:
        test_path = os.path.join(tmpdir, "doc-v2.docx")
        _create_test_docx(test_path)
        parser = DocumentParser()
        doc = parser.parse(test_path)
        ca = build_corpus_analysis([doc])
        scorer = QualityScorer(corpus_analysis=ca)
        card = scorer.score(doc)
        categories = {r.category for r in card.results}
        assert "readability" in categories
        assert "retrieval_aware" in categories


def test_new_weights_sum_to_one():
    """Updated weights (without graph) should sum to ~1.0."""
    from corpus_analyzer import build_corpus_analysis

    with tempfile.TemporaryDirectory() as tmpdir:
        test_path = os.path.join(tmpdir, "doc-v2.docx")
        _create_test_docx(test_path)
        parser = DocumentParser()
        doc = parser.parse(test_path)
        ca = build_corpus_analysis([doc])
        scorer = QualityScorer(corpus_analysis=ca)
        card = scorer.score(doc)
        total_weight = sum(r.weight for r in card.results if r.weight > 0)
        assert abs(total_weight - 1.0) < 0.02, f"Weights sum to {total_weight}"


if __name__ == "__main__":
    test_scoring_detects_expected_issues()
    test_score_breakdown()
    test_new_scoring_criteria_present()
    test_new_weights_sum_to_one()
    print("\nAll tests passed!")
